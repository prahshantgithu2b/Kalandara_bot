#!/usr/bin/env python3
"""
🎯 PERFECT MISSING PERSON FORM BOT
Working, Reliable, and User-Friendly
"""

import logging
from telegram import Update
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes, ConversationHandler
from docx import Document
import re
from datetime import datetime
import os
from config import BOT_TOKEN, DOCUMENT_PATH

# Enable logging
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# States
WAITING_FOR_ANSWER = 1

# 🎯 PERFECT HINGLISH QUESTIONS
QUESTIONS = {
    'SEX': '👤 Missing person ka gender kya hai? (Male/Female/Other)\n💡 Example: Male',
    'NAME W/O HUSBAND NAME': '📝 Missing person ka naam kya hai? (Bina husband ke naam ke)\n💡 Example: Priya Sharma',
    'ADDRESS': '🏠 Missing person ka address kya hai?\n💡 Example: House No. 123, Street 5, Delhi',
    'AGE': '🎂 Missing person ki age kitni hai? (Years mein)\n💡 Example: 25',
    'BUILT': '💪 Missing person ka body type kya hai?\n\n**Options:**\n• Slim/दुबला\n• Medium/मध्यम\n• Healthy/स्वस्थ\n• Heavy/मोटा\n\n💡 Type: Slim, Medium, Healthy, or Heavy',
    'HEIGHT': '📏 Missing person ki height kitni hai?\n💡 Example: 5 feet 6 inches',
    'COMPLEXION': '🎨 Missing person ka complexion kaisa hai?\n\n**Options:**\n• Fair/गोरा\n• Wheatish/गेहूं जैसा\n• Dark/काला\n• Very Fair/बहुत गोरा\n\n💡 Type: Fair, Wheatish, Dark, or Very Fair',
    'HAIR COLOR': '💇‍♀️ Missing person ke baal ka color kya hai?\n\n**Options:**\n• Black/काला\n• Brown/भूरा\n• Grey/सफेद\n• Blonde/सुनहरा\n• Red/लाल\n\n💡 Type: Black, Brown, Grey, Blonde, or Red',
    'FACE': '😊 Missing person ka chehra kaisa hai?\n\n**Options:**\n• Round/गोल\n• Oval/अंडाकार\n• Square/चोकोर\n• Long/लंबा\n• Heart/दिल के आकार का\n\n💡 Type: Round, Oval, Square, Long, or Heart',
    'CLOTHS': '👕 Missing person ne kaunse kapde pehne the?\n💡 Example: Blue shirt, White kurta',
    'FOOTWEAR': '👟 Missing person ne kaunse footwear pehne the?\n💡 Example: Black shoes, White sandals',
    'MISSING DATE': '📅 Kis din se missing hai? (Date)\n💡 Example: 15 July 2024',
    'TIME': '⏰ Missing hone ka time kya tha?\n💡 Example: 2:30 PM',
    'PLACE OF MISSING': '📍 Kahan se missing hua/hui?\n💡 Example: Central Market, Delhi',
    'DD NO.': '📋 DD number kya hai?\n💡 Example: DD-123/2024',
    'REPORT DATE': '📄 Report ki date kya hai?\n💡 Example: 16 July 2024',
    'PS NAME': '🚔 Police station ka naam kya hai?\n💡 Example: Central Police Station',
    'DISTT': '🏛️ District ka naam kya hai?\n💡 Example: New Delhi'
}

# Valid options for dropdown fields
VALID_OPTIONS = {
    'BUILT': ['Slim', 'Medium', 'Healthy', 'Heavy'],
    'COMPLEXION': ['Fair', 'Wheatish', 'Dark', 'Very Fair'],
    'HAIR COLOR': ['Black', 'Brown', 'Grey', 'Blonde', 'Red'],
    'FACE': ['Round', 'Oval', 'Square', 'Long', 'Heart']
}

class DocumentFiller:
    def __init__(self, docx_path):
        self.docx_path = docx_path
        self.placeholders = self.extract_real_placeholders()
        self.user_data = {}
        
    def extract_real_placeholders(self):
        """Extract only meaningful placeholders"""
        doc = Document(self.docx_path)
        placeholders = set()
        
        # Check paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text
            matches = re.findall(r'\{\{([^}]+)\}\}', text)
            placeholders.update(matches)
        
        # Check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        matches = re.findall(r'\{\{([^}]+)\}\}', text)
                        placeholders.update(matches)
        
        # Filter out irrelevant placeholders
        meaningful_placeholders = []
        irrelevant = ['011-27557184', 'Sho-prashantvhr-dl@nic.in']
        
        for placeholder in placeholders:
            cleaned = placeholder.strip()
            cleaned = re.sub(r'^[{\s]+', '', cleaned)
            cleaned = re.sub(r'[}\s]+$', '', cleaned)
            
            if cleaned and cleaned not in irrelevant and cleaned not in meaningful_placeholders:
                meaningful_placeholders.append(cleaned)
        
        return sorted(meaningful_placeholders)
    
    def get_question(self, placeholder):
        """Get question for placeholder"""
        return QUESTIONS.get(placeholder, f'❓ {placeholder} ka value kya hai?')
    
    def replace_placeholders(self, output_path):
        """Replace placeholders in the document and save"""
        doc = Document(self.docx_path)
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, value in self.user_data.items():
                patterns = [
                    f'{{{{{placeholder}}}}}',
                    f'[{placeholder}]',
                    f'{{{placeholder}}}',
                    f'__{placeholder}__'
                ]
                for pattern in patterns:
                    if pattern in paragraph.text:
                        paragraph.text = paragraph.text.replace(pattern, str(value))
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in self.user_data.items():
                            patterns = [
                                f'{{{{{placeholder}}}}}',
                                f'[{placeholder}]',
                                f'{{{placeholder}}}',
                                f'__{placeholder}__'
                            ]
                            for pattern in patterns:
                                if pattern in paragraph.text:
                                    paragraph.text = paragraph.text.replace(pattern, str(value))
        
        doc.save(output_path)
        return output_path

# Global variables
user_sessions = {}
document_filler = DocumentFiller(DOCUMENT_PATH)

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Start the bot"""
    if not update.effective_user or not update.message:
        return ConversationHandler.END
        
    user_id = update.effective_user.id
    user_sessions[user_id] = {
        'current_placeholder_index': 0,
        'data': {}
    }
    
    welcome_message = """
🎯 **Missing Person Form Bot** 🎯

Namaste! 🙏 Main aapko missing person form fill karne mein help karunga.

📋 **Total Questions:** 18
⏱️ **Estimated Time:** 5-10 minutes
🎁 **Result:** Complete filled form

Chaliye shuru karte hain! 🚀

*Type /cancel anytime to stop*
    """
    
    await update.message.reply_text(welcome_message, parse_mode='Markdown')
    await ask_next_question(update, context)

async def ask_next_question(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Ask the next question"""
    if not update.effective_user or not update.message:
        return ConversationHandler.END
        
    user_id = update.effective_user.id
    
    if user_id not in user_sessions:
        await update.message.reply_text("❌ Please start with /start command")
        return ConversationHandler.END
    
    session = user_sessions[user_id]
    current_index = session['current_placeholder_index']
    total_questions = len(document_filler.placeholders)
    
    if current_index >= total_questions:
        await generate_document(update, context)
        return ConversationHandler.END
    
    placeholder = document_filler.placeholders[current_index]
    question = document_filler.get_question(placeholder)
    
    # Create progress bar
    progress_filled = "🟦" * current_index
    progress_empty = "⬜" * (total_questions - current_index)
    progress_bar = progress_filled + progress_empty
    
    message = f"""
📝 **Question {current_index + 1}/{total_questions}**

{question}

{progress_bar}
*Progress: {current_index + 1}/{total_questions}*
    """
    
    if update.message:
        await update.message.reply_text(message, parse_mode='Markdown')

async def handle_answer(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle user's answer"""
    if not update.effective_user or not update.message:
        return ConversationHandler.END
        
    user_id = update.effective_user.id
    
    if user_id not in user_sessions:
        await update.message.reply_text("❌ Please start with /start command")
        return ConversationHandler.END
    
    session = user_sessions[user_id]
    current_index = session['current_placeholder_index']
    
    if current_index >= len(document_filler.placeholders):
        await update.message.reply_text("✅ All questions answered!")
        return ConversationHandler.END
    
    # Get current placeholder and answer
    placeholder = document_filler.placeholders[current_index]
    answer = update.message.text.strip() if update.message and update.message.text else ""
    
    # Validate dropdown selections
    if placeholder in VALID_OPTIONS:
        valid_options = VALID_OPTIONS[placeholder]
        if answer not in valid_options:
            # Show error and ask again
            error_msg = f"❌ **Invalid selection!**\n\nPlease select from these options:\n"
            for option in valid_options:
                error_msg += f"• {option}\n"
            error_msg += "\nTry again:"
            await update.message.reply_text(error_msg, parse_mode='Markdown')
            return WAITING_FOR_ANSWER
    
    # Save the answer
    session['data'][placeholder] = answer
    
    # Move to next question
    session['current_placeholder_index'] += 1
    
    # Show progress
    progress = f"✅ **{placeholder}:** {answer}"
    await update.message.reply_text(progress, parse_mode='Markdown')
    
    # Ask next question
    await ask_next_question(update, context)

async def generate_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Generate the filled document"""
    if not update.effective_user:
        return ConversationHandler.END
        
    user_id = update.effective_user.id
    session = user_sessions[user_id]
    
    # Set the data in document filler
    document_filler.user_data = session['data']
    
    # Generate output filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"filled_form_{timestamp}.docx"
    
    try:
        # Show processing message
        if update.message:
            await update.message.reply_text("🔄 Processing your form... Please wait!")
        
        # Replace placeholders and save
        output_path = document_filler.replace_placeholders(output_filename)
        
        # Send the document
        if update.message:
            with open(output_path, 'rb') as doc:
                await update.message.reply_document(
                    document=doc,
                    filename=output_filename,
                    caption="🎉 **Form Successfully Filled!** 🎉\n\n✅ Aapka form ready hai!\n📥 Download karke use kar sakte hain.\n\nThank you for using our bot! 🙏"
                )
        
        # Clean up
        os.remove(output_path)
        
        # Clear session
        del user_sessions[user_id]
        
    except Exception as e:
        if update.message:
            await update.message.reply_text(f"❌ Error generating document: {str(e)}")

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Cancel the conversation"""
    if not update.effective_user or not update.message:
        return ConversationHandler.END
        
    user_id = update.effective_user.id
    if user_id in user_sessions:
        del user_sessions[user_id]
    
    await update.message.reply_text("❌ Form filling cancelled. Use /start to begin again.")
    return ConversationHandler.END

def main():
    """Start the bot"""
    if BOT_TOKEN == "YOUR_BOT_TOKEN_HERE":
        print("❌ Please set your bot token in config.py!")
        return
    
    # Create application
    application = Application.builder().token(BOT_TOKEN).build()
    
    # Add conversation handler
    conv_handler = ConversationHandler(
        entry_points=[CommandHandler('start', start)],
        states={
            WAITING_FOR_ANSWER: [MessageHandler(filters.TEXT & ~filters.COMMAND, handle_answer)]
        },
        fallbacks=[CommandHandler('cancel', cancel)]
    )
    
    application.add_handler(conv_handler)
    
    # Start the bot
    print("🎯 PERFECT BOT STARTING...")
    print("📱 Bot is ready to receive messages!")
    application.run_polling()

if __name__ == '__main__':
    main() 